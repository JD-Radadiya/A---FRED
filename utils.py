"""
Utility functions for document generation and logging.
"""

import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def setup_logging(log_dir: str = "logs") -> None:
    """
    Set up logging configuration.
    
    Args:
        log_dir: Directory to store log files
    """
    # Create logs directory if it doesn't exist
    os.makedirs(log_dir, exist_ok=True)
    
    # Create log filename with timestamp
    log_filename = os.path.join(
        log_dir,
        f"app_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    )
    
    # Remove any existing handlers to avoid duplicates
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    # Configure logging with INFO level to capture more details
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_filename, mode='a'),
            logging.StreamHandler()
        ],
        force=True
    )
    
    # Set specific loggers to appropriate levels
    logging.getLogger('streamlit').setLevel(logging.WARNING)
    logging.getLogger('urllib3').setLevel(logging.WARNING)
    logging.getLogger('requests').setLevel(logging.WARNING)
    
    logger = logging.getLogger(__name__)
    logger.info(f"Logging initialized. Log file: {log_filename}")
    logger.info(f"Application started at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


def create_word_document(content: str, filename: str) -> str:
    """
    Create a Word document with the given content.
    
    Args:
        content: Text content to add to the document
        filename: Name for the output file (without extension)
        
    Returns:
        Path to the created document
    """
    logger = logging.getLogger(__name__)
    
    try:
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Generated Response', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph()
        timestamp.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        timestamp_run = timestamp.runs[0]
        timestamp_run.font.size = Pt(9)
        timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add separator
        doc.add_paragraph('_' * 80)
        
        # Add content
        content_para = doc.add_paragraph(content)
        content_run = content_para.runs[0]
        content_run.font.size = Pt(11)
        content_run.font.name = 'Calibri'
        
        # Save document
        output_dir = "outputs"
        os.makedirs(output_dir, exist_ok=True)
        
        filepath = os.path.join(output_dir, f"{filename}.docx")
        doc.save(filepath)
        
        logger.info(f"Created Word document: {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        raise


def sanitize_filename(filename: str) -> str:
    """
    Sanitize filename by removing invalid characters.
    
    Args:
        filename: Original filename
        
    Returns:
        Sanitized filename
    """
    # Remove file extension if present
    if '.' in filename:
        filename = filename.rsplit('.', 1)[0]
    
    # Replace invalid characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    # Limit length
    max_length = 100
    if len(filename) > max_length:
        filename = filename[:max_length]
    
    return filename


def format_file_size(size_bytes: int) -> str:
    """
    Format file size in human-readable format.
    
    Args:
        size_bytes: Size in bytes
        
    Returns:
        Formatted size string
    """
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.2f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.2f} TB"
